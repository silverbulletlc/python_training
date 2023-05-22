# pip3 install python-docx 操作word文件需要安装这个库
# pip3 install --upgrade pip 顺便升级pip

def first():
    from docx import Document # 导入库时，名字是docx，不是python-docx
    doc1 = Document()
    doc1.add_heading('如何使用python创建和操作word',0) # 增加标题，0代表最大一级标题
    doc1.save('./create_data/16创建word基本操作.docx')

def start1():
    from docx import Document
    doc1 = Document()
    # 增加标题
    doc1.add_heading('如何使用python创建和操作word',0)
    doc1.add_heading('一级标题',level=1)
    doc1.add_heading('二级标题',level=2)
    doc1.add_heading('三级标题',level=3)
    doc1.add_paragraph('这个是段落')
    doc1.add_paragraph('这个是段落','Title') # 最后的‘Title’等于0级标题
    # 增加段落
    doc1.add_paragraph('我喜欢你沈旭霞，你到底喜欢我吗？你知道吗你真是个小野猫。我喜欢你沈旭霞，你到底喜欢我吗？你知道吗你真是个小野猫。我喜欢你沈旭霞，你到底喜欢我吗？你知道吗你真是个小野猫。')
    paragrah = doc1.add_paragraph('我喜欢你沈旭霞，你到底喜欢我吗？你知道吗你真是个小野猫。我喜欢你沈旭霞，你到底喜欢我吗？你知道吗你真是个小野猫。我喜欢你沈旭霞，你到底喜欢我吗？你知道吗你真是个小野猫。')
    paragrah.add_run('这个是后加的内容，可以紧接着前面不空行')
    # 增加无序列表
    doc1.add_paragraph('哪个你喜欢？')
    doc1.add_paragraph('喜羊羊',style = 'List Bullet')
    doc1.add_paragraph('美羊羊',style = 'List Bullet')
    doc1.add_paragraph('懒羊羊',style = 'List Bullet')
    # 增加有序列表
    doc1.add_paragraph('今年的计划')
    doc1.add_paragraph('学会python',style = 'List Number')
    doc1.add_paragraph('涨工资',style = 'List Number')
    doc1.add_paragraph('找到女朋友',style = 'List Number')
    # 增加引用
    doc1.add_paragraph('这是一个引用',style = 'Intense Quote')
    # 增加图片
    from docx.shared import Inches # 导入英寸单位
    pic = doc1.add_picture('./base_data/美女.jpg',width = Inches(5)) # 可以设置尺寸
    # 使图片适应页面，计算页面宽度和图片宽度的比例sc，高度也对应计算
    height = pic.height
    width = pic.width
    p_width = doc1.sections[0].page_width
    sc = (p_width - doc1.sections[0].left_margin*2)/width # 页面宽度要减去两侧外宽度
    pic.height = int(height * sc)
    pic.width = int(width * sc)
    # 插入表格
    table = doc1.add_table(rows=1,cols=3)
    cells = table.rows[0].cells
    cells[0].text = '编号'
    cells[1].text = '姓名'
    cells[2].text = '职业'
    data = [
        [1,'梁超','天才'],
        [2,'沈旭霞','美女'],
        [3,'沈美芳','妈妈桑'],
    ]
    for i,n,w in data: # 你妹的居然还能这么遍历
        cells = table.add_row().cells
        cells[0].text = str(i) # 数字要转为字符
        cells[1].text = n
        cells[2].text = w

    doc1.save('./create_data/16创建word基本操作.docx')

if __name__ == '__main__':
    # first()
    start1()